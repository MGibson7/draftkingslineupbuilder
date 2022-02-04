import pandas as pd
import pprint
import random
df = pd.read_excel('insertexcelfilehere')


#position, player name, salary, and score

salary = 50000

players = {}
# Get players in a dictionary
for player in range(len(df)):
    a = df.iloc[player, 0]
    b = df.iloc[player, 1]
    c = df.iloc[player, 2]
    d = df.iloc[player, 3]
    e = df.iloc[player, 9]
    f = df.iloc[player, 11]

    players[player] = {'Position': a, 'Name': b, 'Salary':c, 'Team':d, 'Score':e, 'Opponent':f}

#print(players)
#print(players[0]['Position'])

line_ups_scores = []
line_up = []
line_up_score = []
line_up_scoresss = []
#0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0
line_up_score_top_20 = [0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0]
line_ups = {}
wr_te_teams = []
rb_wr_te_teams = []

line_ups_generated = -2

def make_lineup(total_line_ups):
  line_up = []
  x = 0
  line_up_positions = []
  line_up_salary = []
  line_up_score = []
  up = 0
  line_ups_generated = -1
  wr_te_teams = []
  rb_wr_te_teams = []
  while line_ups_generated < total_line_ups:
      #3rd lowest salary hb
      third_lowest_salary_hb = 4700
      number_of_defenses = 20
      line_up = []
      line_up_salary = []
      line_up_score = []
      line_up_positions = []
      # Change the below numbers where QB starts, RB starts, etc where they are on excel - 2 #RB Starts subtract one from actual start should be same as QB stop
      rb_starts = 36
      numofplayers = 141
      qb_starts = 20
      qb_stops = 36
      x = random.randint(qb_starts, qb_stops)
      line_ups_generated = line_ups_generated + 1
      print(f"line_ups_generated {line_ups_generated}")
      while len(line_up) <9:
        salary_left = 50000-sum(line_up_salary)-2000
        if x >= len(df)-2:
            # Change the below range based on where HB starts and WR stops
            x = random.randint(rb_starts,numofplayers)
        elif 'QB' not in line_up_positions:
            # Change the below range based on where QB starts and stops
            x = random.randint(qb_starts, qb_stops)
            if players[x]['Position'] == 'QB':
                line_up.append(players[x]['Name'])
                line_up_salary.append(players[x]['Salary'])
                line_up_score.append(players[x]['Score'])
                line_up_positions.append('QB')
                qb_team = (players[x]['Team'])
                qb_opponent_team = (players[x]['Opponent'])
                # Change the below range based on where HB starts and WR stops
                x = random.randint(rb_starts, numofplayers)
            else:
                x = x
        elif players[x]['Name'] in line_up:
            x = x
        elif players[x]['Salary'] > salary_left:
            x = x
        elif salary_left - (players[x]['Salary']) < (8-len(line_up)-1)*(third_lowest_salary_hb):
            x = x
        elif players[x]['Position'] == 'RB':
            if line_up_positions.count('RB')<2:
                line_up.append(players[x]['Name'])
                line_up_salary.append(players[x]['Salary'])
                line_up_score.append(players[x]['Score'])
                line_up_positions.append('RB')
                rb_wr_te_teams.append((players[x]['Team']))
                # Change the below range based on where HB starts and WR stops
                x = random.randint(rb_starts, numofplayers)
            elif line_up_positions.count('RB') == 2 and line_up_positions.count('WR')<= 3:
                line_up.append(players[x]['Name'])
                line_up_salary.append(players[x]['Salary'])
                line_up_score.append(players[x]['Score'])
                line_up_positions.append('RB')
                rb_wr_te_teams.append((players[x]['Team']))
                # Change the below range based on where HB starts and WR stops
                x = random.randint(rb_starts, numofplayers)
        elif players[x]['Position'] == 'WR':
            if line_up_positions.count('WR')<3:
                line_up.append(players[x]['Name'])
                line_up_salary.append(players[x]['Salary'])
                line_up_score.append(players[x]['Score'])
                line_up_positions.append('WR')
                rb_wr_te_teams.append((players[x]['Team']))
                wr_te_teams.append((players[x]['Team']))
                # Change the below range based on where HB starts and WR stops
                x = random.randint(rb_starts, numofplayers)
            elif line_up_positions.count('RB') <= 2 and line_up_positions.count('WR')== 3:
                line_up.append(players[x]['Name'])
                line_up_salary.append(players[x]['Salary'])
                line_up_score.append(players[x]['Score'])
                line_up_positions.append('WR')
                rb_wr_te_teams.append((players[x]['Team']))
                wr_te_teams.append((players[x]['Team']))
                # Change the below range based on where HB starts and WR stops
                x = random.randint(rb_starts, numofplayers)
        elif players[x]['Position'] == 'TE':
            if 'TE' not in line_up_positions:
                line_up.append(players[x]['Name'])
                line_up_salary.append(players[x]['Salary'])
                line_up_score.append(players[x]['Score'])
                line_up_positions.append('TE')
                rb_wr_te_teams.append((players[x]['Team']))
                wr_te_teams.append((players[x]['Team']))
                # Change the below range based on where HB starts and WR stops
                x = random.randint(rb_starts, numofplayers)
            elif line_up_positions.count('RB') < 3 and line_up_positions.count('WR') < 4:
                if line_up_positions.count('TE')<1:
                    line_up.append(players[x]['Name'])
                    line_up_salary.append(players[x]['Salary'])
                    line_up_score.append(players[x]['Score'])
                    line_up_positions.append('TE')
                    rb_wr_te_teams.append((players[x]['Team']))
                    wr_te_teams.append((players[x]['Team']))
                    # Change the below range based on where HB starts and WR stops
                    x = random.randint(rb_starts, numofplayers)
        elif players[x]['Position'] == 'DST':
            x = random.randint(rb_starts, numofplayers)
        x = x +1
        if len(line_up) == 8:
            #print(x)
            #a = sum(line_up_salary)
            #print(f"{line_up} {a}")
            d = 0
            line_up.append(df.iloc[0,1])
            line_up_salary.append(df.iloc[0,2])
            #print(df.iloc[0,2])
            line_up_positions.append(df.iloc[0,0])
            line_up_score.append(df.iloc[0,9])
            #print(line_up)
            line_up_3 = sorted(line_up)
            #print(line_up_3)
            if qb_team not in wr_te_teams:
                line_up = []
                line_up_salary = []
                line_up_score = []
                line_up_positions = []
                wr_te_teams = []
                rb_wr_te_teams = []
            #DOUBLE STACK
            elif wr_te_teams.count(qb_team) != 2:
                line_up = []
                line_up_salary = []
                line_up_score = []
                line_up_positions = []
                wr_te_teams = []
                rb_wr_te_teams = []
            elif qb_opponent_team not in rb_wr_te_teams:
                line_up = []
                line_up_salary = []
                line_up_score = []
                line_up_positions = []
                wr_te_teams = []
                rb_wr_te_teams = []
            else:
                while d <number_of_defenses:
                    #print(sum(line_up_salary))
                    if (sum(line_up_salary) <= 50000) and (sum(line_up_salary) >= 49000):
                        if 0==0:
                            print(f"New Line up! {line_up_3} {sum(line_up_salary)}")
                            up = up + 1
                            line_up_2 = line_up_3.copy()
                            line_up_score_2 = line_up_score.copy()
                            line_up_scoresss.append(sum(line_up_score_2))
                            line_ups['line_up' + str(up)] = line_up_2
                            line_ups['line_up_score' + str(up)] = sum(line_up_score_2)
                            if min(line_up_score_top_20) > sum(line_up_score_2):
                                line_up_scoresss.pop(-1)
                                #print(line_up_scoresss)
                            else:
                                q = min(line_up_score_top_20)
                                if q in line_up_scoresss:
                                    line_up_scoresss.remove(q)
                                line_up_score_top_20.remove(q)
                                line_up_score_top_20.append(sum(line_up_score_2))
                            line_up.pop(-1)
                            line_up_salary.pop(-1)
                            line_up_score.pop(-1)
                            line_up_positions.pop(-1)
                            #print(line_up)
                            d = d +1
                            line_up.append(df.iloc[d, 1])
                            line_up_salary.append(df.iloc[d, 2])
                            line_up_positions.append(df.iloc[d, 0])
                            line_up_score.append(df.iloc[d, 9])
                            #print(line_up)
                            line_up_3 = sorted(line_up)
                            #print(line_up_3)
                        else:
                            d = d + 1
                            line_up.pop(-1)
                            line_up_salary.pop(-1)
                            line_up_score.pop(-1)
                            line_up_positions.pop(-1)
                            line_up.append(df.iloc[d, 1])
                            line_up_salary.append(df.iloc[d, 2])
                            line_up_positions.append(df.iloc[d, 0])
                            line_up_score.append(df.iloc[d, 9])
                            line_up_3 = sorted(line_up)
                    else:
                        d = d +1
                        line_up.pop(-1)
                        line_up_salary.pop(-1)
                        line_up_score.pop(-1)
                        line_up_positions.pop(-1)
                        line_up.append(df.iloc[d, 1])
                        line_up_salary.append(df.iloc[d, 2])
                        line_up_positions.append(df.iloc[d, 0])
                        line_up_score.append(df.iloc[d, 9])
                        line_up_3 = sorted(line_up)
                line_up_salary.append(50000)
                line_up = ['Calvin Johnson', 'Bo Jackson', 'Randy Moss', 'Favre', 'Owens', 'Manning', 'Harrison',
                           'Johnson', 'Clausen', 'Grier']
                wr_te_teams = []
                rb_wr_te_teams = []






make_lineup(1000)
#pprint.pprint(line_ups)


#j = pprint.pprint(line_ups)

line_up_scoresss.sort(reverse = True)

p =0
l = 1
import re
import docx
from docx import Document
from docx.shared import Inches

document = Document()
dup_line_ups = []

for number in range(len(line_up_scoresss)):
    z = str(line_up_scoresss[p])
    for a_key in line_ups.keys():
        print('BUILDING...')
        r = str(line_ups[a_key])
        if (r == z):
            #print(a_key)
            a_key_2 = re.findall(r'\d', a_key)
            m = 0
            a_key_3 = ''
            for n in range (len(a_key_2)):
                a_key_3 = a_key_3 + str(a_key_2[m])
                m = m + 1
            a_key_4 = 'line_up' + a_key_3
            h = line_ups[a_key_4]
            if h not in dup_line_ups:
                print("HARD")
                dup_line_ups.append(h)
                document.add_paragraph(f"Line Up #{l}")
                l = l + 1
                document.add_paragraph(f"{line_ups[a_key_4]}")
                document.add_paragraph(f"Total Score {z}")
                print("Added a line up")
    p = p +1
    document.save(f'draftkingslineups.docx')


g = int((len(line_ups.keys())))
print(g)
h = int(g/2)
print(h)

print(line_up_scoresss)



from datetime import datetime

now = datetime.now()

current_time = now.strftime("%H:%M:%S")
print("Current Time =", current_time)
